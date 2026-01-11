import os
import json
import asyncio
import logging
import aiofiles
from pathlib import Path
from typing import List, Tuple
from datetime import datetime
import streamlit as st
import sqlite3
from dotenv import load_dotenv
from openai import OpenAI, APIError
from docx import Document

from duckduckgo_search import DDGS
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import warnings
import fitz

warnings.filterwarnings("ignore")
load_dotenv()

# CONFIG
CONFIG = {
    'app_name': 'AI Document Assistant Pro v2.0',
    'max_context_tokens': 4000,
    'chunk_size': 150,
    'chunk_overlap': 30,
    'top_k': 3,
    'embedding_model': 'paraphrase-multilingual-MiniLM-L12-v2',
    'openrouter_model': 'meta-llama/llama-3.2-3b-instruct:free',
    'documents_dir': 'documents',
    'answers_dir': 'answers',
    'db_path': 'knowledgebase_pro.db',
    'retry_attempts': 3,
    'use_web_search_default': True,
    'min_doc_context_chars': 500,
    'min_doc_chunks': 1,
    'web_max_results': 3
}

# Создание директорий
os.makedirs(CONFIG['documents_dir'], exist_ok=True)
os.makedirs(CONFIG['answers_dir'], exist_ok=True)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


@st.cache_resource
def load_embedding_model():
    st.info('🔄 Загрузка модели эмбеддингов (~30 сек)...')
    model = SentenceTransformer(CONFIG['embedding_model'], device='cpu')
    st.success('✅ Модель готова!')
    return model


class OpenRouterClient:
    def __init__(self, model_id: str = None):
        self.api_key = os.getenv('OPENAI_API_KEY')
        self.model_id = model_id or CONFIG['openrouter_model']
        if not self.api_key:
            st.error("❌ OPENAI_API_KEY не найден! Создайте .env файл")
        self.client = OpenAI(api_key=self.api_key, base_url="https://openrouter.ai/api/v1")

    async def ask(self, question: str, context: str) -> str:
        if not self.api_key:
            return "❌ Настройте OPENAI_API_KEY в .env"

        try:
            messages = [
                {"role": "system",
                 "content": "Ты эксперт по анализу документов. Используй только предоставленный контекст."},
                {"role": "user", "content": f"Контекст:\n{context}\n\nВопрос: {question}"}
            ]
            response = await asyncio.to_thread(
                self.client.chat.completions.create,
                model=self.model_id,
                messages=messages,
                max_tokens=CONFIG['max_context_tokens']
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"❌ Ошибка API: {str(e)}"


class VectorKnowledgeBase:
    def __init__(self, db_path: str):
        self.db_path = db_path
        self.model = load_embedding_model()
        self.init_db()
        self.chunks, self.embeddings = self.load_chunks()

    def init_db(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT,
                chunk_text TEXT,
                chunk_index INTEGER,
                doc_type TEXT,
                embedding BLOB
            )
        """)
        conn.commit()
        conn.close()

    def chunk_text(self, text: str) -> List[str]:
        if len(text) < 200:  # Маленький текст целиком
            return [text] if len(text) > 50 else []

        # 1. По словам (основной метод)
        words = text.split()
        chunks = []
        step = max(1, CONFIG['chunk_size'] - CONFIG['chunk_overlap'])  # Минимум 1!

        for i in range(0, len(words), step):
            chunk = ' '.join(words[i:i + CONFIG['chunk_size']])
            if len(chunk) > 50:
                chunks.append(chunk)

        # 2. Fallback: по символам если слов мало
        if len(chunks) == 0 and len(text) > 200:
            for i in range(0, len(text), 800):
                chunk = text[i:i + 800].strip()
                if len(chunk) > 100:
                    chunks.append(chunk)

        print(f"🔍 CHUNKING: {len(text.split())} слов → {len(chunks)} чанков")
        return chunks[:50]  # Максимум 50 чанков на документ

    def add_document(self, filename: str, content: str, doc_type: str):
        # ✅ Фиксируем имя файла правильно
        if not filename or len(filename) > 100:
            filename = "document.pdf"

        clean_filename = Path(filename).stem[:50] + f".{doc_type}"  # Kolobkov.pdf
        print(f"🔍 SAVE: '{clean_filename}' → {len(content)} символов → {len(self.chunk_text(content))} чанков")

        chunks = self.chunk_text(content)
        if not chunks:
            print(f"❌ ПУСТЫЕ ЧАНКИ для {clean_filename}")
            return

        embeddings = self.model.encode(chunks)
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        for i, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            cursor.execute("""
                INSERT OR REPLACE INTO chunks (filename, chunk_text, chunk_index, doc_type, embedding)
                VALUES (?, ?, ?, ?, ?)
            """, (clean_filename, chunk, i, doc_type, emb.tobytes()))
        conn.commit()
        conn.close()
        self.chunks, self.embeddings = self.load_chunks()
        st.success(f"✅ {clean_filename}: {len(chunks)} чанков")

    def load_chunks(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT chunk_text, embedding FROM chunks")
        chunks = []
        embeddings = []
        for chunk_text, emb_bytes in cursor.fetchall():
            chunks.append(chunk_text)
            embeddings.append(np.frombuffer(emb_bytes, dtype=np.float32))
        conn.close()
        return chunks, np.array(embeddings) if embeddings else np.array([])

    def search(self, query: str, top_k: int = CONFIG['top_k']) -> str:
        if len(self.chunks) == 0:
            return "База знаний пуста. Загрузите документы."

        query_emb = self.model.encode([query])
        similarities = cosine_similarity(query_emb, self.embeddings)[0]
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        context = '\n'.join([self.chunks[i] for i in top_indices])
        return context[:CONFIG['max_context_tokens']]


async def process_document(filepath: str) -> Tuple[str, str]:
    filename = Path(filepath).name
    content = ""

    if filepath.lower().endswith('.txt'):
        async with aiofiles.open(filepath, 'r', encoding='utf-8') as f:
            content = await f.read()
        doc_type = 'txt'
        logging.info(f"TXT {filename}: {len(content)} символов")

    elif filepath.lower().endswith('.pdf'):
        try:
            doc = fitz.open(filepath)
            pages_text = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                pages_text.append(f"📄 Страница {page_num + 1}:\n{text}")  # ✅ ВСЕ страницы!

            content = '\n\n---\n\n'.join(pages_text)
            #doc.close()
            doc_type = 'pdf'
            logging.info(f"PDF {filename}: {len(doc)} страниц, {len(content)} символов")


        except Exception as e:
            content = f"Ошибка PDF: {e}"
            doc_type = 'pdf_error'

    elif filepath.lower().endswith('.docx'):
        try:
            doc = Document(filepath)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = '\n\n'.join(paragraphs) if paragraphs else "Пустой DOCX"
            doc_type = 'docx'
            logging.info(f"DOCX {filename}: {len(paragraphs) if paragraphs else 0} параграфов")
        except Exception as e:
            content = f"Ошибка DOCX: {e}"
            doc_type = 'docx_error'

    else:
        content = "Поддержка: TXT/PDF/DOCX"
        doc_type = 'unsupported'

    return content, doc_type


def save_answer_docx(question: str, answer: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]
    filename = Path(CONFIG['answers_dir']) / f"answer_{timestamp}.docx"
    doc = Document()
    doc.add_heading(CONFIG['app_name'], level=1)
    doc.add_heading(f"Вопрос: {question}", level=2)
    doc.add_paragraph(f"Ответ:\n{answer}")
    doc.save(filename)
    return str(filename)


def web_search(query: str, max_results: int = CONFIG['web_max_results']) -> str:
    try:
        with DDGS() as ddgs:
            results = [r.get('body') or r.get('snippet', '')
                       for r in ddgs.text(query, max_results=max_results)]
        return '\n'.join(results)
    except:
        return ""


async def ask_ai(client, kb, question: str) -> str:
    use_web = st.session_state.get('use_web_search', True)
    doc_context = kb.search(question)

    # Веб-поиск если мало контекста
    min_chars = st.session_state.get('min_doc_context_chars', 500)
    few_chunks = len(kb.chunks) < 2
    short_context = len(doc_context) < min_chars

    web_context = ""
    if use_web and (few_chunks or short_context):
        web_context = web_search(question)

    full_context = f"{doc_context}\n\n{'Доп. веб: ' + web_context if web_context else ''}"
    return await client.ask(question, full_context)


def main():
    st.set_page_config(page_title=CONFIG['app_name'], layout='wide', initial_sidebar_state='expanded')
    st.title(CONFIG['app_name'])
    st.markdown("---")

    # Инициализация
    if 'client' not in st.session_state:
        st.session_state.client = OpenRouterClient()
    if 'kb' not in st.session_state:
        st.session_state.kb = VectorKnowledgeBase(CONFIG['db_path'])

    client = st.session_state.client
    kb = st.session_state.kb

    with st.spinner('🔄 Инициализация...'):
        st.success(f"✅ Готово! Чанков: {len(kb.chunks)}")
        # Sidebar статистика
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📊 Чанков", len(kb.chunks))
        with col2:
            unique_files = len(
                set(chunk.split('filename=')[1].split(',')[0] for chunk in kb.chunks if 'filename=' in chunk))
            st.metric("📄 Файлов", unique_files)

        if st.button("📋 Показать файлы"):
            files = list(set(Path(c).parent.name for c in kb.chunks))
            st.write(f"**Файлы в базе:** {', '.join(files)}")

    # Sidebar
    with st.sidebar:
        st.header("📁 Документы")
        uploaded_files = st.file_uploader(
            "Загрузите документы",
            accept_multiple_files=True,
            type=['txt', 'pdf', 'docx']
        )

        if uploaded_files:
            for file in uploaded_files:
                filepath = Path(CONFIG['documents_dir']) / file.name
                with open(filepath, 'wb') as f:
                    f.write(file.getbuffer())
                try:
                    content, doc_type = asyncio.run(process_document(str(filepath)))
                    kb.add_document(file.name, content, doc_type)
                    st.session_state.kb = VectorKnowledgeBase(CONFIG['db_path'])  # Обновить сессию
                except Exception as e:
                    st.error(f"❌ {file.name}: {e}")

        st.header("🤖 Модель ИИ")
        model_options = [
        "allenai/molmo-2-8b:free",
        "bytedance-seed/seedream-4.5",
        "xiaomi/mimo-v2-flash:free",
        "mistralai/devstral-2512:free",
        "sourceful/riverflow-v2-max-preview",
        "meta-llama/llama-3.2-3b-instruct",  # ✅ БАЗОВАЯ - всегда работает
        ]

        selected_model = st.selectbox("Выберите:", model_options, key="model")
        if selected_model != client.model_id:
            st.session_state.client = OpenRouterClient(selected_model)
            st.rerun()

        st.header("⚙️ Настройки")
        st.session_state.use_web_search = st.checkbox("🔍 Веб-поиск", value=True)
        st.session_state.min_doc_context_chars = st.number_input("Мин. символов",
                                                                 value=CONFIG['min_doc_context_chars'], min_value=0,
                                                                 step=100)

        if st.button("🗑️ Очистить базу"):
            conn = sqlite3.connect(CONFIG['db_path'])
            conn.execute("DELETE FROM chunks")
            conn.commit()
            conn.close()
            st.session_state.kb = VectorKnowledgeBase(CONFIG['db_path'])
            st.success("✅ База очищена!")
            st.rerun()

    # Главный интерфейс
    col1, col2 = st.columns([3, 1])
    with col1:
        question = st.text_area("❓ Вопрос по документам:", height=100, key="question")
    with col2:
        if st.button("🚀 Спросить ИИ", type="primary") and question:
            with st.spinner('🤖 Генерация...'):
                answer = asyncio.run(ask_ai(client, kb, question))
                st.markdown(f"**Ответ:**\n{answer}")
                st.session_state.last_answer = answer

    if st.session_state.get('last_answer') and st.button("💾 Сохранить DOCX"):
        path = save_answer_docx(st.session_state.last_question or question,
                                st.session_state.last_answer)
        st.success(f"✅ {path}")



if __name__ == "__main__":
    main()
